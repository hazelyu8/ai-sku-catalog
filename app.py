import streamlit as st
import sqlite3
import pytesseract
from PIL import Image
from docx import Document
import pandas as pd
import openpyxl
import os
import shutil

# ===============================================================
# 1. OCR & DOCX Extraction
# ===============================================================
def extract_text_from_image(image_file):
    """Extract text from an uploaded image using OCR."""
    img = Image.open(image_file)
    text = pytesseract.image_to_string(img)
    return text

def extract_text_from_docx(docx_file):
    """Extract both paragraph and table text from DOCX."""
    doc = Document(docx_file)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) == 2:
                text_parts.append(f"{cells[0]} {cells[1]}")
            elif cells:
                text_parts.append(" | ".join(cells))

    return "\n".join(text_parts)

# ===============================================================
# 2. Parse Fields
# ===============================================================
def parse_vendor_doc(text):
    """Parse structured fields from vendor document text."""
    fields = {
        "Product Description": None,
        "Batch/Lot No.": None,
        "Date": None,
        "SKU": None,
        "Qty": None,
    }

    for line in text.splitlines():
        line = line.strip()
        for key in fields.keys():
            if line.lower().startswith(key.lower()):
                parts = line.split(":", 1)
                if len(parts) == 2:
                    fields[key] = parts[1].strip()
    return fields

# ===============================================================
# 3. Database Operations
# ===============================================================
def save_to_db(fields):
    """Save parsed data to local SQLite database — skip duplicates."""
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS sku_catalog (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        product_desc TEXT,
        batch_lot TEXT,
        date TEXT,
        sku TEXT,
        qty TEXT
    )
    """)

    sku = fields["SKU"]
    batch = fields["Batch/Lot No."]

    # Check for duplicates
    cur.execute("SELECT COUNT(*) FROM sku_catalog WHERE sku = ? AND batch_lot = ?", (sku, batch))
    exists = cur.fetchone()[0]

    if exists > 0:
        st.warning(f"⚠️ Entry with SKU '{sku}' and Batch '{batch}' already exists. Skipped saving.")
    else:
        cur.execute("""
        INSERT INTO sku_catalog (product_desc, batch_lot, date, sku, qty)
        VALUES (?, ?, ?, ?, ?)
        """, (
            fields["Product Description"],
            fields["Batch/Lot No."],
            fields["Date"],
            fields["SKU"],
            fields["Qty"]
        ))
        conn.commit()
        st.success(f"✅ Saved new entry: SKU {sku}, Batch {batch}")
    conn.close()

def clear_database():
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute("DELETE FROM sku_catalog;")
    conn.commit()
    conn.close()

# ===============================================================
# 4. Save to Excel
# ===============================================================
def save_to_excel(
    fields,
    excel_path="data/specs/sample_specs.xlsx",
    sheet_name="Master Sheet - 12th Floor"
):
    """Write parsed data into Excel sheet starting at row 757."""
    import openpyxl

    if not os.path.exists(excel_path):
        st.error(f"❌ Excel file not found at: {excel_path}")
        return

    temp_copy = excel_path.replace(".xlsx", "_temp.xlsx")
    shutil.copyfile(excel_path, temp_copy)

    try:
        wb = openpyxl.load_workbook(temp_copy)
    except Exception as e:
        st.error(f"❌ Could not open Excel file. Make sure it's a valid .xlsx and not open in Excel.\n\nError: {e}")
        return

    if sheet_name not in wb.sheetnames:
        st.error(f"❌ Sheet '{sheet_name}' not found. Available sheets: {wb.sheetnames}")
        return

    ws = wb[sheet_name]
    target_row = 757

    # Fill empty rows up to 757
    if ws.max_row < target_row:
        for _ in range(ws.max_row, target_row - 1):
            ws.append([])

    # Insert data (no Customer)
    data = [
        fields.get("Product Description"),
        fields.get("Batch/Lot No."),
        fields.get("Date"),
        fields.get("SKU"),
        fields.get("Qty"),
    ]

    for col, value in enumerate(data, start=1):
        ws.cell(row=target_row, column=col, value=value)

    wb.save(temp_copy)
    shutil.move(temp_copy, excel_path)
    st.info(f"✅ Data inserted into row {target_row} of '{sheet_name}'")

# ===============================================================
# 5. Streamlit UI
# ===============================================================
st.title("🤖 AI SKU Agent — Offline Chat Mode")

st.markdown("""
Welcome to your **AI-Powered SKU Catalog Agent**!  
You can either upload a document manually below **or chat** with your agent:
- Type **"process the new SKU doc"** to process the latest file in `data/specs/`
- Type **"show database"** to display current entries
- Type **"clear database"** to reset all records
""")

# ------------------- Chat Memory -------------------
if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

user_input = st.chat_input("Type a message (e.g., 'process the new SKU doc')")

# ------------------- Chat Agent Logic -------------------
if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input})
    response = None

    if "process" in user_input.lower():
        folder = "data/specs/"
        files = [os.path.join(folder, f) for f in os.listdir(folder) if f.endswith((".docx", ".jpg", ".png", ".jpeg"))]
        if not files:
            response = "❌ No files found in data/specs/."
        else:
            latest = max(files, key=os.path.getmtime)
            text = extract_text_from_docx(latest) if latest.endswith(".docx") else extract_text_from_image(latest)
            fields = parse_vendor_doc(text)
            save_to_db(fields)
            save_to_excel(fields)
            response = f"✅ Added SKU {fields.get('SKU')} (Batch {fields.get('Batch/Lot No.')}) to Master Sheet."

    elif "show" in user_input.lower() and "database" in user_input.lower():
        conn = sqlite3.connect("sku_catalog.db")
        df = pd.read_sql_query("SELECT * FROM sku_catalog", conn)
        conn.close()
        st.dataframe(df)
        response = "📊 Here's your current database."

    elif "clear" in user_input.lower() and "database" in user_input.lower():
        clear_database()
        response = "🧹 Database cleared successfully."

    else:
        response = "🤖 I can process new SKU docs, show the database, or clear it."

    st.session_state.messages.append({"role": "assistant", "content": response})

# ------------------- Manual Upload Mode -------------------
st.divider()
st.header("📄 Manual Upload Mode")

uploaded_file = st.file_uploader("Upload a vendor document (JPG, PNG, DOCX)", type=["jpg", "png", "jpeg", "docx"])

if uploaded_file:
    st.success("✅ File uploaded successfully!")

    text = extract_text_from_docx(uploaded_file) if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" else extract_text_from_image(uploaded_file)
    st.subheader("🔎 Extracted Text")
    st.text(text if text.strip() else "(No text detected)")

    fields = parse_vendor_doc(text)
    st.subheader("📦 Parsed Fields")
    st.json(fields)

    if st.button("💾 Save to Database & Excel"):
        if any(fields.values()):
            save_to_db(fields)
            save_to_excel(fields)
            st.success("✅ Data saved successfully to both Database and Excel!")
        else:
            st.warning("⚠️ No valid fields detected — please check your document format.")
