import streamlit as st
import sqlite3
import pytesseract
from PIL import Image
from docx import Document
import re
import pandas as pd
import openpyxl
import os

# ------------------------
# 1. OCR & DOCX Extraction
# ------------------------
def extract_text_from_image(image_file):
    """Extract text from an uploaded image using OCR."""
    img = Image.open(image_file)
    text = pytesseract.image_to_string(img)
    return text

def extract_text_from_docx(docx_file):
    """Extract text from .docx files, including tables."""
    doc = Document(docx_file)
    text = []

    # Regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    # Text inside tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text.append(" | ".join(cells))

    return "\n".join(text)

# ------------------------
# 2. Parse Fields (Regex)
# ------------------------
def parse_vendor_doc(text):
    """Parse structured fields (Customer, SKU, Qty, etc.) from the extracted text."""
    fields = {
        "Customer": None,
        "Product Description": None,
        "Batch/Lot No.": None,
        "Date": None,
        "SKU": None,
        "Qty": None
    }

    # Handles formats like "Customer | Tarte" or "Customer: Tarte"
    patterns = {
        "Customer": r"Customer[:\|\-]?\s*\n?\s*([A-Za-z0-9 \-]+)",
        "Product Description": r"(?:Product Description|Description)[:\|\-]?\s*\n?\s*([A-Za-z0-9 \-]+)",
        "Batch/Lot No.": r"(?:Batch|Lot|Batch/Lot No.)[:\|\-]?\s*\n?\s*([A-Za-z0-9\-\_]+)",
        "Date": r"Date[:\|\-]?\s*\n?\s*([0-9/]+)",
        "SKU": r"SKU[:\|\-]?\s*\n?\s*([A-Za-z0-9\-\_]+)",
        "Qty": r"(?:Qty|Quantity)[:\|\-]?\s*\n?\s*([0-9]+)"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            fields[key] = match.group(1).strip()

    return fields

# ------------------------
# 3. Save to SQLite
# ------------------------
def save_to_db(fields):
    """Save parsed fields into a SQLite database."""
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS sku_catalog (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        customer TEXT,
        product_desc TEXT,
        batch_lot TEXT,
        date TEXT,
        sku TEXT,
        qty TEXT
    )
    """)

    cur.execute("""
    INSERT INTO sku_catalog (customer, product_desc, batch_lot, date, sku, qty)
    VALUES (?, ?, ?, ?, ?, ?)
    """, (
        fields["Customer"], fields["Product Description"], fields["Batch/Lot No."],
        fields["Date"], fields["SKU"], fields["Qty"]
    ))

    conn.commit()
    conn.close()

# ------------------------
# 4. Save to Excel
# ------------------------
def save_to_excel(fields, excel_path="data/specs/sample_specs.xlsx", sheet_name="Master Sheet - 12th Floor"):
    """Save parsed fields into the Excel sheet."""
    os.makedirs(os.path.dirname(excel_path), exist_ok=True)

    if not os.path.exists(excel_path):
        # Create a new Excel file with headers
        df = pd.DataFrame(columns=["Customer", "Product Description", "Batch/Lot No.", "Date", "SKU", "Qty"])
        df.to_excel(excel_path, index=False, sheet_name=sheet_name)

    wb = openpyxl.load_workbook(excel_path)

    if sheet_name not in wb.sheetnames:
        ws = wb.create_sheet(sheet_name)
        ws.append(["Customer", "Product Description", "Batch/Lot No.", "Date", "SKU", "Qty"])
    else:
        ws = wb[sheet_name]

    ws.append([
        fields["Customer"], fields["Product Description"], fields["Batch/Lot No."],
        fields["Date"], fields["SKU"], fields["Qty"]
    ])

    wb.save(excel_path)

# ------------------------
# 5. Streamlit Interface
# ------------------------
st.title("📦 AI-Powered SKU Catalog Agent")
st.markdown("Upload a vendor document (.docx or image). The app will extract product info and automatically save it to Excel and the local database.")

uploaded_file = st.file_uploader("Upload a vendor document (JPG, PNG, DOCX)", type=["jpg", "png", "jpeg", "docx"])

if uploaded_file:
    st.success("✅ File uploaded successfully!")

    # Extract text
    if uploaded_file.type in ["image/jpeg", "image/png"]:
        text = extract_text_from_image(uploaded_file)
    else:
        text = extract_text_from_docx(uploaded_file)

    st.subheader("🔎 Extracted Text")
    st.text(text if text.strip() else "(No text detected)")

    # Parse text into structured fields
    fields = parse_vendor_doc(text)
    st.subheader("📦 Parsed Fields")
    st.json(fields)

    # Save button
    if st.button("💾 Save to Database & Excel"):
        if any(fields.values()):
            save_to_db(fields)
            save_to_excel(fields)
            st.success("✅ Entry saved to database and Excel successfully!")
        else:
            st.warning("⚠️ No valid fields detected. Please check your document format.")
