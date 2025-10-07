from docx import Document

path = "data/specs/smolder_eyes_document.docx"
doc = Document(path)

print("📄 INSPECTING DOCUMENT STRUCTURE\n")

# 1️⃣ Paragraphs
print("Paragraphs:")
for i, para in enumerate(doc.paragraphs):
    print(f"  [{i}] {repr(para.text)}")

# 2️⃣ Tables
print("\nTables:")
for t_idx, table in enumerate(doc.tables):
    print(f"  Table {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        row_texts = [cell.text.strip() for cell in row.cells]
        print(f"    Row {r_idx}: {row_texts}")

print("\n✅ Done inspecting document.")

